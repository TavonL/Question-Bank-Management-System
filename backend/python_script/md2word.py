import pypandoc
import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
def convertFile(filename, src_type="md", tgt_type="docx", extra_args=['-f','markdown']):
    '''
    文件类型转换，传入paperName
    '''
    try:
        import pypandoc
        src_name = filename + '.' + src_type
        tgt_name = filename + '.' + tgt_type
        pypandoc.convert_file(src_name, to=tgt_type, format=src_type, outputfile=tgt_name, extra_args=extra_args)
        return 0
    except Exception as e:
        print(e)
        return -1

def addHeader(filename):
    '''
    设置docx文档标题，传入paperName
    '''
    document = Document(filename + '.docx')
    subtitle_text = "学号 _________    姓名 _________    成绩 _________"
    subtitle = document.paragraphs[0].insert_paragraph_before(text = subtitle_text)
    subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title = document.paragraphs[0].insert_paragraph_before(text = filename)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.save(filename + '.docx')

def createMdFile(filename, content):
    '''
    传入paperName和questionsContent
    '''

    with open(filename+'.md', 'w') as f:
        f.write(content)

def deleteFiles(filename):
    '''
    用于删除上述md和docx文件，传入paperName
    '''
    os.remove(filename + '.md')
    os.remove(filename + '.docx')

if __name__ == '__main__':
    filename = '2011年卷'
    # createMdFile(filename, 'testesteststestet')
    
    convertFile(filename)
    addHeader(filename)

    # deleteFiles(filename)
    