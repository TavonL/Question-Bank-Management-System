import pypandoc
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
def convert_file(filename, src_type="md", tgt_type="docx"):
    try:
        import pypandoc
        src_name = filename + '.' + src_type
        tgt_name = filename + '.' + tgt_type
        pypandoc.convert_file(src_name, to=tgt_type, format=src_type, outputfile=tgt_name)
        return 0
    except Exception as e:
        print(e)
        return -1

def addHeader(filename):
    document = Document(filename + '.docx')
    subtitle_text = "学号 _________    姓名 _________    成绩 _________"
    subtitle = document.paragraphs[0].insert_paragraph_before(text = subtitle_text)
    subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title = document.paragraphs[0].insert_paragraph_before(text = filename)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.save(filename + '.docx')
if __name__ == '__main__':
    filename = '上海大学卷'
    convert_file(filename)
    addHeader(filename)